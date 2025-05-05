import os
import io
from django.shortcuts import render, get_object_or_404
from django.core.files.base import ContentFile
from django.http import JsonResponse, FileResponse, HttpResponse, Http404
from django.conf import settings
from .models import Report
from docx import Document
from reportlab.pdfgen import canvas


def view_reports(request):
    reports = Report.objects.all().order_by('-created_at')
    return render(request, 'reports/list.html', {'reports': reports})


def generate_docx(request, report_id):
    report = get_object_or_404(Report, id=report_id)

    # Create a Word document in memory
    doc = Document()
    doc.add_heading(report.title, level=1)
   
    doc.add_paragraph(f"Created at: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")  # Blank line
    doc.add_paragraph(report.content)

    # Save document to in-memory buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Save the generated docx to the model field
    filename = f"{report.title.replace(' ', '_')}.docx"
    report.docx_file.save(filename, ContentFile(docx_buffer.read()))
    report.save()

    return HttpResponse("Docx generated successfully.")


def download_docx(request, report_id):
    report = get_object_or_404(Report, id=report_id)

    if not report.docx_file:
        return HttpResponse("No DOCX file available for this report.", status=404)

    try:
        return FileResponse(open(report.docx_file.path, 'rb'), as_attachment=True)
    except FileNotFoundError:
        raise Http404("DOCX file not found on the server.")


def generate_pdf(request, report_id):
    report = get_object_or_404(Report, id=report_id)

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer)

    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, 800, report.title)

    p.setFont("Helvetica", 12)
  
    p.drawString(100, 760, f"Created at: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    
    text_object = p.beginText(100, 740)
    text_object.setFont("Helvetica", 11)
    for line in report.content.splitlines():
        text_object.textLine(line)
    p.drawText(text_object)

    p.showPage()
    p.save()
    buffer.seek(0)

    filename = f"{report.title.replace(' ', '_')}.pdf"
    report.pdf_file.save(filename, ContentFile(buffer.read()))
    report.save()

    return HttpResponse("PDF generated successfully.")


def download_pdf(request, report_id):
    report = get_object_or_404(Report, id=report_id)

    if not report.pdf_file:
        return HttpResponse("No PDF file available for this report.", status=404)

    try:
        return FileResponse(open(report.pdf_file.path, 'rb'), as_attachment=True)
    except FileNotFoundError:
        raise Http404("PDF file not found on the server.")
